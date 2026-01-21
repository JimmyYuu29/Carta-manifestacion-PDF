"""
Word Renderer - Document rendering with docxtpl
Renderizador de Word con docxtpl
"""

from pathlib import Path
from typing import List, Tuple, Optional
import re
from copy import deepcopy

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .plugin_loader import PluginPack
from .context_builder import ContextBuilder
from .rule_engine import RuleEngine, EvaluationTrace


class DocxRenderer:
    """
    Word document renderer using python-docx
    Renderizador de documentos Word usando python-docx
    """

    def __init__(self, plugin: PluginPack):
        self.plugin = plugin
        self.context_builder = ContextBuilder(plugin)
        self.rule_engine = RuleEngine(plugin)
        self._template_path: Optional[Path] = None

    def render(self, data: dict, output_path: Path, template_path: Optional[Path] = None) -> Tuple[Path, List[EvaluationTrace]]:
        """
        Render Word document
        Renderizar documento Word

        Args:
            data: Input data dictionary
            output_path: Path for output file
            template_path: Optional custom template path

        Returns:
            Tuple of (output_path, evaluation_traces)
        """
        # 1. Load template
        if template_path:
            self._template_path = template_path
        else:
            self._template_path = self.plugin.get_template_path()

        doc = Document(self._template_path)

        # 2. Build context
        context = self.context_builder.build_context(data)

        # 3. Get conditional values (si/no)
        conditionals = self.context_builder.get_conditional_values(data)
        context.update(conditionals)

        # 4. Evaluate rules
        visibility_map, traces = self.rule_engine.evaluate_all_rules(data)
        context["visibility"] = visibility_map

        # 5. Strip conditional blocks
        self._strip_conditional_blocks(doc, conditionals)

        # 6. Replace variables
        self._process_document(doc, context, conditionals)

        # 7. Post-process
        self._post_process(doc)

        # 8. Save
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

        return output_path, traces

    def _strip_conditional_blocks(self, doc: Document, cond_values: dict) -> None:
        """
        Remove content between {% if VAR == 'si' %} ... {% endif %}
        when cond_values[VAR] == 'no'
        """
        body_elems = list(doc.element.body)
        inside_remove = False
        inside_keep = False
        trash = []

        for el in body_elems:
            txt = ""
            if el.tag.endswith('p'):
                txt = "".join(t.text or "" for t in el.iter() if getattr(t, "text", None)).strip()

            # Opening block
            m_open = re.match(r"\{% if (\w+)\s*==\s*'si' %\}", txt)
            if m_open:
                var = m_open.group(1)
                if cond_values.get(var, 'no') == 'si':
                    inside_keep = True
                else:
                    inside_remove = True
                trash.append(el)
                continue

            # Closing block
            if re.match(r"\{% endif %\}", txt):
                trash.append(el)
                inside_remove = False
                inside_keep = False
                continue

            if inside_remove:
                trash.append(el)

        for el in trash:
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)

    def _process_document(self, doc: Document, context: dict, conditionals: dict) -> None:
        """Process all paragraphs, tables, headers and footers / Procesar todos los parrafos, tablas, encabezados y pies de pagina"""
        # Process paragraphs
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            if original_text.strip():
                new_text = self._replace_variables(original_text, context, conditionals)
                if new_text != original_text:
                    original_format = self._save_paragraph_format(paragraph)
                    paragraph.clear()
                    paragraph.text = new_text
                    self._restore_paragraph_format(paragraph, original_format)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        original_text = paragraph.text
                        if original_text.strip():
                            new_text = self._replace_variables(original_text, context, conditionals)
                            if new_text != original_text:
                                paragraph.text = new_text

        # Process headers and footers in all sections
        for section in doc.sections:
            # Process header
            if section.header:
                for paragraph in section.header.paragraphs:
                    original_text = paragraph.text
                    if original_text.strip():
                        new_text = self._replace_variables(original_text, context, conditionals)
                        if new_text != original_text:
                            paragraph.text = new_text
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                original_text = paragraph.text
                                if original_text.strip():
                                    new_text = self._replace_variables(original_text, context, conditionals)
                                    if new_text != original_text:
                                        paragraph.text = new_text

            # Process footer
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    original_text = paragraph.text
                    if original_text.strip():
                        new_text = self._replace_variables(original_text, context, conditionals)
                        if new_text != original_text:
                            paragraph.text = new_text
                for table in section.footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                original_text = paragraph.text
                                if original_text.strip():
                                    new_text = self._replace_variables(original_text, context, conditionals)
                                    if new_text != original_text:
                                        paragraph.text = new_text

    def _replace_variables(self, text: str, variables: dict, conditionals: dict) -> str:
        """Replace variables and process conditionals / Reemplazar variables y procesar condicionales"""
        # Process inline conditionals
        text = self._process_conditionals(text, conditionals)

        # Handle lista_alto_directores special pattern
        lista_pattern = r'\{\{lista_alto_directores:[^}]+\}\}'
        lista_matches = list(re.finditer(lista_pattern, text, re.DOTALL))

        for match in reversed(lista_matches):
            if 'lista_alto_directores' in variables and variables['lista_alto_directores']:
                text = text[:match.start()] + variables['lista_alto_directores'] + text[match.end():]
            else:
                text = text[:match.start()] + text[match.end():]

        # Replace simple variables
        for var_name, var_value in variables.items():
            if var_name == 'lista_alto_directores':
                continue
            if var_name in ('visibility', 'texts'):
                continue

            patterns = [
                rf'\{{\{{\s*{re.escape(var_name)}\s*\}}\}}',
                rf'\{{\{{\s*{re.escape(var_name)}\s*\|\s*int\s*\}}\}}',
                rf'\{{\{{\s*{re.escape(var_name)}\s*\|\s*int\s*-\s*1\s*\}}\}}'
            ]

            for pattern in patterns:
                if '|int - 1' in pattern and var_value:
                    try:
                        replacement = str(int(var_value) - 1)
                    except (ValueError, TypeError):
                        replacement = str(var_value) if var_value else ''
                else:
                    replacement = str(var_value) if var_value else ''

                text = re.sub(pattern, replacement, text)

        # Clean remaining markers
        text = re.sub(r'\[?\{\{[^}]*\}\}\]?', '', text)
        text = re.sub(r'\[\]\.mark', '', text)
        text = re.sub(r'\.mark', '', text)
        text = re.sub(r'\[\.mark\]', '', text)

        return text

    def _process_conditionals(self, text: str, conditionals: dict) -> str:
        """Process conditional blocks / Procesar bloques condicionales"""
        for cond_var, cond_value in conditionals.items():
            # Pattern with mark
            if_pattern = rf'\[\{{% if {cond_var} == \'si\' %\}}\]\.mark(.*?)\[\{{% endif %\}}\]\.mark'
            if cond_value == 'si':
                text = re.sub(if_pattern, r'\1', text, flags=re.DOTALL)
            else:
                text = re.sub(if_pattern, '', text, flags=re.DOTALL)

            # Pattern without mark
            if_pattern = rf'\{{% if {cond_var} == \'si\' %\}}(.*?)\{{% endif %\}}'
            if cond_value == 'si':
                text = re.sub(if_pattern, r'\1', text, flags=re.DOTALL)
            else:
                text = re.sub(if_pattern, '', text, flags=re.DOTALL)

        # Clean remaining conditional markers
        text = re.sub(r'\{%[^%]*%\}', '', text)

        return text

    def _save_paragraph_format(self, paragraph) -> dict:
        """Save paragraph formatting / Guardar formato de parrafo"""
        format_info = {
            'alignment': paragraph.alignment,
            'style': paragraph.style.name if paragraph.style else None,
            'runs': []
        }

        for run in paragraph.runs:
            run_format = {
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'font_color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
            }
            format_info['runs'].append(run_format)

        return format_info

    def _restore_paragraph_format(self, paragraph, format_info: dict) -> None:
        """Restore paragraph formatting / Restaurar formato de parrafo"""
        if format_info['alignment']:
            paragraph.alignment = format_info['alignment']

        if format_info['style']:
            try:
                paragraph.style = format_info['style']
            except Exception:
                pass

        if format_info['runs'] and paragraph.runs:
            for i, run in enumerate(paragraph.runs):
                if i < len(format_info['runs']):
                    run_format = format_info['runs'][i]
                    if run_format['bold'] is not None:
                        run.bold = run_format['bold']
                    if run_format['italic'] is not None:
                        run.italic = run_format['italic']
                    if run_format['underline'] is not None:
                        run.underline = run_format['underline']
                    if run_format['font_name']:
                        run.font.name = run_format['font_name']
                    if run_format['font_size']:
                        run.font.size = run_format['font_size']

    def _post_process(self, doc: Document) -> None:
        """Post-processing: cell coloring, remove empty paragraphs, fix numbering"""
        self._apply_cell_colors(doc)
        self._remove_underlines(doc)
        self._fix_numbering(doc)

    def _apply_cell_colors(self, doc: Document) -> None:
        """Apply cell colors based on content / Aplicar colores a celdas"""
        colors = self.plugin.formatting.get("colors", {})

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip().lower()
                    if text in colors:
                        self._set_cell_color(cell, colors[text])

    def _set_cell_color(self, cell, hex_color: str) -> None:
        """Set cell background color / Establecer color de fondo de celda"""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), hex_color.replace("#", ""))
        tcPr.append(shd)

    def _remove_underlines(self, doc: Document) -> None:
        """Remove underlines from all runs / Quitar subrayados de todos los runs"""
        for p in doc.paragraphs:
            for run in p.runs:
                run.underline = False

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.underline = False

    def _fix_numbering(self, doc: Document) -> None:
        """Fix paragraph numbering / Corregir numeracion de parrafos"""
        current_number = 1
        sub_number = 1
        in_sub_list = False

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()

            # Main points (start with number followed by dot)
            main_match = re.match(r'^(\d+)\.\s+(.+)', text)
            if main_match:
                paragraph.text = f"{current_number}. {main_match.group(2)}"
                current_number += 1
                in_sub_list = False

            # Sub-points (start with lowercase letter followed by dot)
            sub_match = re.match(r'^[a-z]\.\s+(.+)', text)
            if sub_match:
                if not in_sub_list:
                    sub_number = 1
                    in_sub_list = True

                letter = chr(ord('a') + sub_number - 1)
                paragraph.text = f"{letter}. {sub_match.group(1)}"
                sub_number += 1
